# /// script
# dependencies = [
#   "python-docx>=1.1.0",
#   "jsonschema>=4.20.0",
# ]
# requires-python = ">=3.10"
# ///

"""
AIS Proposal Document Generator

Reads a JSON input file, validates it against the proposal schema, opens the
AIS-branded template.docx, performs placeholder replacement while preserving
all existing styles, generates body content sections, injects comments, and
writes the output .docx file.

Usage:
    uv run scripts/generate.py --input data.json --output proposal.docx
    uv run scripts/generate.py --input data.json --output proposal.docx --template assets/template.docx
"""

import argparse
import json
import os
import re
import sys
import copy
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from xml.etree import ElementTree as ET

import jsonschema
from docx import Document
from docx.oxml.ns import qn
from lxml import etree as lxml_etree

SKILL_DIR = Path(__file__).resolve().parent.parent
DEFAULT_TEMPLATE = SKILL_DIR / "assets" / "template.docx"
SCHEMA_PATH = SKILL_DIR / "assets" / "proposal-schema.json"

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

# Maps content block types from JSON schema to Word style names
CONTENT_STYLE_MAP = {
    "body": "_Body",
    "body_0_after": "_Body_0_after",
    "heading_a": "_Heading A",
    "heading_b": "_Heading B",
    "bullet1": "_Bullet1",
    "bullet1_0_after": "_Bullet1_0_after",
    "bullet2": "_Bullet2",
    "bullet2_0_after": "_Bullet2_0_after",
    "bullet3": "_Bullet3",
    "bullet3_0_after": "_Bullet3_0_after",
    "callout_title": "_Callout Title",
    "callout_text": "_Callout Text",
    "callout_bullet": "_Callout Bullet",
}

DEFAULT_DISCLAIMER = (
    "This proposal includes data that shall not be disclosed outside the Government "
    "and shall not be duplicated, used, or disclosed-in whole or in part-for any "
    "purpose other than to evaluate this proposal. If, however, a contract is awarded "
    "to this offeror as a result of-or in connection with-the submission of this data, "
    "the Government shall have the right to duplicate, use, or disclose the data to "
    "the extent provided in the resulting contract. This restriction does not limit "
    "the Government's right to use information contained in this data if it is obtained "
    "from another source without restriction. The data subject to this restriction are "
    'contained on sheets with the following disclosure legend: "Use or disclosure of '
    "information contained on this sheet is subject to the restriction on the title "
    'page of this proposal."'
)


def load_and_validate(input_path: str) -> dict:
    """Load JSON input and validate against the proposal schema."""
    with open(input_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    with open(SCHEMA_PATH, "r", encoding="utf-8") as f:
        schema = json.load(f)

    jsonschema.validate(instance=data, schema=schema)
    return data


def replace_text_in_runs(paragraph, replacements: dict):
    """Replace placeholder text across runs in a paragraph.

    Handles the common OOXML case where {{placeholder}} is split across
    multiple <w:r> elements. Joins all run text, performs replacement,
    then redistributes text back to runs preserving formatting.
    """
    runs = paragraph.runs
    if not runs:
        return

    full_text = "".join(r.text or "" for r in runs)

    changed = False
    for old, new in replacements.items():
        if old in full_text:
            full_text = full_text.replace(old, new)
            changed = True

    if not changed:
        return

    # Redistribute text across runs: put all text in first run, clear the rest.
    # This preserves the first run's formatting for the entire paragraph.
    if runs:
        runs[0].text = full_text
        for r in runs[1:]:
            r.text = ""


def replace_in_xml_part(part_xml: bytes, replacements: dict) -> bytes:
    """Replace placeholder text in raw XML bytes (for headers/footers)."""
    # Register namespaces to avoid ns0: prefixes
    for prefix, uri in [
        ("w", NS_W),
        ("r", NS_R),
        ("mc", "http://schemas.openxmlformats.org/markup-compatibility/2006"),
        ("w14", "http://schemas.microsoft.com/office/word/2010/wordml"),
        ("w15", "http://schemas.microsoft.com/office/word/2012/wordml"),
        ("wp", "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"),
        ("a", "http://schemas.openxmlformats.org/drawingml/2006/main"),
        ("pic", "http://schemas.openxmlformats.org/drawingml/2006/picture"),
    ]:
        ET.register_namespace(prefix, uri)

    root = ET.fromstring(part_xml)

    for t_elem in root.findall(f".//{{{NS_W}}}t"):
        if t_elem.text:
            for old, new in replacements.items():
                if old in t_elem.text:
                    t_elem.text = t_elem.text.replace(old, new)

    return ET.tostring(root, encoding="unicode", xml_declaration=True).encode("utf-8")


def build_replacements(data: dict) -> dict:
    """Build a flat replacement map from the structured input data."""
    cp = data["cover_page"]
    sb = cp["submitted_by"]
    st = cp["submitted_to"]

    replacements = {
        "{{title}}": cp["title"],
        "{{solicitation_number}}": cp["solicitation_number"],
        "{{date}}": cp["date"],
        "{{volume_name}}": cp["volume_name"],
        "{{volume_factor}}": cp["volume_factor"],
        "{{company_name}}": sb["name"],
        "{{company_address_line1}}": sb["address"]["line1"],
        "{{company_address_line2}}": sb["address"].get("line2", ""),
        "{{company_city_state_zip}}": sb["address"]["city_state_zip"],
        "{{uei}}": sb.get("uei", ""),
        "{{contract_contact_name}}": sb["contact"]["name"],
        "{{contract_contact_email}}": sb["contact"]["email"],
        "{{contract_contact_phone}}": sb["contact"].get("phone", ""),
        "{{agency_name}}": st["agency_name"],
        "{{agency_org}}": st["org_name"],
        "{{agency_address_line1}}": st["address"]["line1"],
        "{{agency_city_state_zip}}": st["address"]["city_state_zip"],
    }

    if "contracting_officer" in st:
        co = st["contracting_officer"]
        replacements["{{agency_co_name}}"] = co["name"]
        replacements["{{agency_co_email}}"] = co["email"]
    if "contract_specialist" in st:
        cs = st["contract_specialist"]
        replacements["{{agency_cs_name}}"] = cs["name"]
        replacements["{{agency_cs_email}}"] = cs["email"]

    # Header
    replacements["{{header_text}}"] = data["header"]["text"]

    # Footer
    footer = data.get("footer", {})
    replacements["{{footer_disclaimer}}"] = footer.get(
        "disclaimer_text",
        "Use or disclosure of data contained on this sheet is subject to the "
        "restriction on the title page of this proposal.",
    )

    # Disclaimer
    replacements["{{disclaimer_text}}"] = cp.get("disclaimer", DEFAULT_DISCLAIMER)

    return replacements


def find_body_section_start(doc: Document) -> int:
    """Find the index of the first Heading1 paragraph in the body (after section break)."""
    for i, para in enumerate(doc.paragraphs):
        if para.style and para.style.name == "Heading 1":
            return i
    return -1


def clear_body_content(doc: Document, start_idx: int):
    """Remove all body content paragraphs from start_idx to the end."""
    body = doc.element.body
    paras = list(body.findall(f"{{{NS_W}}}p"))

    # Find the actual XML elements for paragraphs at and after start_idx
    doc_paras = doc.paragraphs
    if start_idx >= len(doc_paras):
        return

    elements_to_remove = []
    for i in range(start_idx, len(doc_paras)):
        elements_to_remove.append(doc_paras[i]._element)

    for elem in elements_to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)


def add_section_content(doc: Document, sections: list, insert_before_element=None):
    """Add body content sections with proper styles.

    Each section gets a Heading1 paragraph followed by content blocks.
    """
    body = doc.element.body

    # Find the final sectPr element — we need to insert before it
    sect_pr = body.find(f"{{{NS_W}}}sectPr")

    for section in sections:
        # Add Heading1
        heading_para = doc.add_paragraph(section["heading"])
        heading_para.style = doc.styles["Heading 1"]
        if sect_pr is not None:
            body.remove(heading_para._element)
            sect_pr.addprevious(heading_para._element)

        # Add content blocks
        for block in section["content"]:
            style_name = CONTENT_STYLE_MAP.get(block["type"], "_Body")

            if block.get("bold_label"):
                # Create paragraph with bold label + normal text
                para = doc.add_paragraph()
                para.style = doc.styles[style_name]
                bold_run = para.add_run(block["bold_label"] + " ")
                bold_run.bold = True
                para.add_run(block["text"])
            else:
                para = doc.add_paragraph(block["text"])
                para.style = doc.styles[style_name]

            if sect_pr is not None:
                body.remove(para._element)
                sect_pr.addprevious(para._element)


def inject_comments(doc: Document, comments: list, sections: list):
    """Inject comments into the document, anchored to section headings.

    Comments are stored in word/comments.xml and referenced from document.xml
    via commentRangeStart/End and commentReference elements.
    """
    if not comments:
        return

    # Find Heading1 paragraphs in the document
    heading_paras = []
    for para in doc.paragraphs:
        if para.style and para.style.name == "Heading 1":
            heading_paras.append(para)

    # Access the document's part to get/create comments.xml
    doc_part = doc.part

    # Build comments XML using lxml (python-docx uses lxml internally)
    nsmap = {"w": NS_W, "r": NS_R}
    comments_root = lxml_etree.Element(qn("w:comments"), nsmap=nsmap)

    comment_id_start = 100  # Start from 100 to avoid conflicts

    for i, comment_data in enumerate(comments):
        section_idx = comment_data["section_index"]
        if section_idx >= len(heading_paras):
            print(
                f"Warning: comment references section {section_idx} but only "
                f"{len(heading_paras)} sections exist. Skipping.",
                file=sys.stderr,
            )
            continue

        cid = str(comment_id_start + i)
        author = comment_data.get("author", "AIS Proposal Generator")
        now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:00Z")

        # Create comment element
        comment_el = lxml_etree.SubElement(comments_root, qn("w:comment"))
        comment_el.set(qn("w:id"), cid)
        comment_el.set(qn("w:author"), author)
        comment_el.set(qn("w:date"), now)

        # Comment paragraph
        cp = lxml_etree.SubElement(comment_el, qn("w:p"))
        cr = lxml_etree.SubElement(cp, qn("w:r"))
        ct = lxml_etree.SubElement(cr, qn("w:t"))
        ct.text = comment_data["text"]

        # Add comment anchors to the heading paragraph (lxml elements)
        para_elem = heading_paras[section_idx]._element

        # Insert commentRangeStart before the first run
        range_start = lxml_etree.Element(qn("w:commentRangeStart"))
        range_start.set(qn("w:id"), cid)
        para_elem.insert(0, range_start)

        # Insert commentRangeEnd and commentReference after the last run
        range_end = lxml_etree.SubElement(para_elem, qn("w:commentRangeEnd"))
        range_end.set(qn("w:id"), cid)

        ref_run = lxml_etree.SubElement(para_elem, qn("w:r"))
        ref_rpr = lxml_etree.SubElement(ref_run, qn("w:rPr"))
        ref_style = lxml_etree.SubElement(ref_rpr, qn("w:rStyle"))
        ref_style.set(qn("w:val"), "CommentReference")
        ref_ref = lxml_etree.SubElement(ref_run, qn("w:commentReference"))
        ref_ref.set(qn("w:id"), cid)

    return comments_root


def _fix_mc_ignorable_namespaces(doc_xml_bytes: bytes, ns_map: dict) -> bytes:
    """Ensure every prefix listed in mc:Ignorable has a matching xmlns declaration.

    python-docx/lxml strips namespace declarations that aren't directly used in
    elements, but mc:Ignorable requires them to be declared. Word shows
    'unreachable content' if any listed prefix lacks an xmlns declaration.
    """
    import re

    doc_str = doc_xml_bytes.decode("utf-8")

    ignorable_match = re.search(r'mc:Ignorable="([^"]*)"', doc_str)
    if not ignorable_match:
        return doc_xml_bytes

    prefixes = ignorable_match.group(1).split()
    insertions = []
    for prefix in prefixes:
        if prefix in ns_map and f'xmlns:{prefix}=' not in doc_str:
            insertions.append(f'xmlns:{prefix}="{ns_map[prefix]}"')

    if not insertions:
        return doc_xml_bytes

    insert_str = " " + " ".join(insertions)
    doc_str = doc_str.replace(
        f'mc:Ignorable="{ignorable_match.group(1)}"',
        insert_str + f' mc:Ignorable="{ignorable_match.group(1)}"',
    )

    return doc_str.encode("utf-8")


def validate_docx(path: str, strict: bool = False) -> list[dict]:
    """Validate a .docx file for common OOXML issues that cause Word errors.

    Returns a list of issue dicts: {severity, part, message}.
    Severity is 'error' (will cause Word errors) or 'warning' (may cause issues).
    If strict=True, warnings are promoted to errors.
    """
    issues = []

    def _add(severity: str, part: str, msg: str):
        issues.append({"severity": severity, "part": part, "message": msg})

    if not os.path.exists(path):
        _add("error", path, "File does not exist")
        return issues

    try:
        zf = zipfile.ZipFile(path, "r")
    except zipfile.BadZipFile:
        _add("error", path, "Not a valid ZIP/DOCX file")
        return issues

    with zf:
        names = set(zf.namelist())

        # --- Check 1: Required parts exist ---
        for required in ["[Content_Types].xml", "word/document.xml"]:
            if required not in names:
                _add("error", required, "Required part is missing")

        # --- Check 2: XML well-formedness ---
        for name in sorted(names):
            if not name.endswith(".xml") and not name.endswith(".rels"):
                continue
            try:
                raw = zf.read(name)
                ET.fromstring(raw)
            except ET.ParseError as e:
                _add("error", name, f"Malformed XML: {e}")

        # --- Check 3: mc:Ignorable namespace declarations ---
        mc_ns_map = {
            'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
            'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
            'w16': 'http://schemas.microsoft.com/office/word/2018/wordml',
            'w16cex': 'http://schemas.microsoft.com/office/word/2018/wordml/cex',
            'w16cid': 'http://schemas.microsoft.com/office/word/2016/wordml/cid',
            'w16du': 'http://schemas.microsoft.com/office/word/2023/wordml/word16du',
            'w16sdtdh': 'http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash',
            'w16sdtfl': 'http://schemas.microsoft.com/office/word/2021/wordml/sdtformatlock',
            'w16se': 'http://schemas.microsoft.com/office/word/2015/wordml/symex',
            'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing14',
        }
        for name in sorted(names):
            if not name.endswith(".xml"):
                continue
            content = zf.read(name).decode("utf-8", errors="replace")
            ig_match = re.search(r'mc:Ignorable="([^"]*)"', content)
            if not ig_match:
                continue
            for prefix in ig_match.group(1).split():
                if f"xmlns:{prefix}=" not in content:
                    _add("error", name,
                         f'mc:Ignorable lists "{prefix}" but xmlns:{prefix} is not declared')

        # --- Check 4: Content_Types coverage ---
        if "[Content_Types].xml" in names:
            ct_raw = zf.read("[Content_Types].xml").decode("utf-8", errors="replace")
            for name in sorted(names):
                if name.startswith("word/") and name.endswith(".xml"):
                    part_name = "/" + name
                    if part_name not in ct_raw and name != "word/_rels/document.xml.rels":
                        ext = name.rsplit(".", 1)[-1]
                        ext_default = f'Extension="{ext}"'
                        if ext_default not in ct_raw:
                            _add("warning", name, "No Content_Types entry (Override or Default)")

        # --- Check 5: Relationship targets resolve ---
        for name in sorted(names):
            if not name.endswith(".rels"):
                continue
            try:
                rels_raw = zf.read(name)
                rels_tree = ET.fromstring(rels_raw)
            except ET.ParseError:
                continue
            rels_dir = os.path.dirname(name)
            base_dir = os.path.dirname(rels_dir) if rels_dir.endswith("_rels") else rels_dir
            for rel in rels_tree:
                target = rel.get("Target", "")
                target_mode = rel.get("TargetMode", "")
                if target_mode == "External" or target.startswith("http"):
                    continue
                resolved = os.path.normpath(os.path.join(base_dir, target)).replace("\\", "/")
                if resolved not in names:
                    _add("warning", name,
                         f'Relationship {rel.get("Id")} targets "{target}" '
                         f'(resolved: {resolved}) which is missing from the archive')

        # --- Check 6: Comment integrity ---
        if "word/document.xml" in names:
            doc_content = zf.read("word/document.xml").decode("utf-8", errors="replace")
            comment_refs = set(re.findall(r'w:commentRangeStart\s[^>]*w:id="(\d+)"', doc_content))
            comment_ends = set(re.findall(r'w:commentRangeEnd\s[^>]*w:id="(\d+)"', doc_content))
            comment_references = set(re.findall(r'w:commentReference\s[^>]*w:id="(\d+)"', doc_content))

            if comment_refs:
                if "word/comments.xml" not in names:
                    _add("error", "word/comments.xml",
                         f"Document has {len(comment_refs)} comment anchors but no comments.xml")
                else:
                    comments_content = zf.read("word/comments.xml").decode("utf-8", errors="replace")
                    defined_ids = set(re.findall(r'w:id="(\d+)"', comments_content))
                    orphaned = comment_refs - defined_ids
                    if orphaned:
                        _add("error", "word/comments.xml",
                             f"Comment IDs {orphaned} referenced in document but not defined")

                for cid in comment_refs:
                    if cid not in comment_ends:
                        _add("error", "word/document.xml",
                             f"commentRangeStart id={cid} has no matching commentRangeEnd")
                    if cid not in comment_references:
                        _add("warning", "word/document.xml",
                             f"commentRangeStart id={cid} has no matching commentReference")

        # --- Check 7: Empty .rels files ---
        for name in sorted(names):
            if not name.endswith(".rels"):
                continue
            try:
                rels_raw = zf.read(name)
                rels_tree = ET.fromstring(rels_raw)
                children = list(rels_tree)
                if len(children) == 0:
                    _add("warning", name, "Empty .rels file (no Relationship elements)")
            except ET.ParseError:
                pass

        # --- Check 8: Residual personal/machine data ---
        sensitive_patterns = [
            (r'[A-Z]:\\Users\\[^"<]+', "Local file path detected"),
            (r'file:///[A-Z]:', "file:// URI with drive letter detected"),
        ]
        skip_parts = {"word/fontTable.xml"}
        for name in sorted(names):
            if not name.endswith(".xml") or name in skip_parts:
                continue
            content = zf.read(name).decode("utf-8", errors="replace")
            for pattern, msg in sensitive_patterns:
                matches = re.findall(pattern, content)
                if matches:
                    _add("warning", name, f"{msg}: {matches[0]}")

    # Promote warnings to errors in strict mode
    if strict:
        for issue in issues:
            if issue["severity"] == "warning":
                issue["severity"] = "error"

    return issues


def print_validation_results(issues: list[dict], path: str) -> bool:
    """Print validation results. Returns True if no errors."""
    errors = [i for i in issues if i["severity"] == "error"]
    warnings = [i for i in issues if i["severity"] == "warning"]

    if not issues:
        print(f"✅ {path}: All checks passed")
        return True

    if errors:
        print(f"❌ {path}: {len(errors)} error(s), {len(warnings)} warning(s)")
    else:
        print(f"⚠️  {path}: {len(warnings)} warning(s)")

    for issue in issues:
        icon = "❌" if issue["severity"] == "error" else "⚠️ "
        print(f"  {icon} [{issue['part']}] {issue['message']}")

    return len(errors) == 0


def process_template(template_path: str, data: dict, output_path: str):
    """Main processing pipeline: template → filled document."""
    replacements = build_replacements(data)

    # Step 1: Raw XML replacement for headers/footers (not accessible via python-docx)
    import tempfile

    temp_path = output_path + ".tmp"

    header_footer_files = [
        "word/header1.xml",
        "word/header2.xml",
        "word/header3.xml",
        "word/footer1.xml",
        "word/footer2.xml",
        "word/footer3.xml",
    ]

    with zipfile.ZipFile(template_path, "r") as zin:
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                raw = zin.read(item.filename)
                if item.filename in header_footer_files:
                    raw = replace_in_xml_part(raw, replacements)
                zout.writestr(item, raw)

    # Step 2: Open with python-docx for document body manipulation
    doc = Document(temp_path)

    # Step 3: Replace placeholders in all paragraphs (cover page + TOC + body)
    for para in doc.paragraphs:
        replace_text_in_runs(para, replacements)

    # Also replace in SDT (structured document tag) content
    body_elem = doc.element.body
    for sdt in body_elem.findall(f".//{{{NS_W}}}sdt"):
        for para_elem in sdt.findall(f".//{{{NS_W}}}p"):
            from docx.text.paragraph import Paragraph

            para = Paragraph(para_elem, doc.element.body)
            replace_text_in_runs(para, replacements)

    # Step 4: Clear existing body content and insert new sections
    sections = data.get("sections", [])
    if sections:
        # Find and remove existing body Heading1+Body paragraphs
        start_idx = find_body_section_start(doc)
        if start_idx >= 0:
            clear_body_content(doc, start_idx)

        # Add new section content
        add_section_content(doc, sections)

    # Step 5: Inject comments
    comments_data = data.get("comments", [])
    comments_xml_root = None
    if comments_data:
        comments_xml_root = inject_comments(doc, comments_data, sections)

    # Step 6: Save the document
    doc.save(output_path)

    # Step 7: Post-process the zip to fix comments and namespace declarations.
    # python-docx/lxml strips namespace declarations that mc:Ignorable references,
    # which causes Word to show "unreachable content" warnings.
    MC_IGNORABLE_NAMESPACES = {
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
        'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
        'w16': 'http://schemas.microsoft.com/office/word/2018/wordml',
        'w16cex': 'http://schemas.microsoft.com/office/word/2018/wordml/cex',
        'w16cid': 'http://schemas.microsoft.com/office/word/2016/wordml/cid',
        'w16du': 'http://schemas.microsoft.com/office/word/2023/wordml/word16du',
        'w16sdtdh': 'http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash',
        'w16sdtfl': 'http://schemas.microsoft.com/office/word/2021/wordml/sdtformatlock',
        'w16se': 'http://schemas.microsoft.com/office/word/2015/wordml/symex',
        'wp14': 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing14',
    }

    has_comments = comments_xml_root is not None
    comments_xml_bytes = None
    if has_comments:
        comments_xml_bytes = lxml_etree.tostring(
            comments_xml_root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    COMMENTS_CT = (
        'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
    )
    COMMENTS_REL_TYPE = (
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
    )

    final_temp = output_path + ".final.tmp"
    with zipfile.ZipFile(output_path, "r") as zin:
        with zipfile.ZipFile(final_temp, "w", zipfile.ZIP_DEFLATED) as zout:
            has_comments_xml = False
            for item in zin.infolist():
                raw = zin.read(item.filename)

                if item.filename == "word/comments.xml" and has_comments:
                    zout.writestr(item, comments_xml_bytes)
                    has_comments_xml = True
                elif item.filename.endswith('.xml') and b'mc:Ignorable' in raw:
                    raw = _fix_mc_ignorable_namespaces(raw, MC_IGNORABLE_NAMESPACES)
                    zout.writestr(item, raw)
                elif item.filename == "[Content_Types].xml" and has_comments:
                    ct_tree = lxml_etree.fromstring(raw)
                    ct_ns = 'http://schemas.openxmlformats.org/package/2006/content-types'
                    existing = ct_tree.findall(
                        f'{{{ct_ns}}}Override[@PartName="/word/comments.xml"]'
                    )
                    if not existing:
                        lxml_etree.SubElement(ct_tree, f'{{{ct_ns}}}Override', {
                            'PartName': '/word/comments.xml',
                            'ContentType': COMMENTS_CT,
                        })
                    zout.writestr(item, lxml_etree.tostring(
                        ct_tree, xml_declaration=True, encoding="UTF-8", standalone=True
                    ))
                elif item.filename == "word/_rels/document.xml.rels" and has_comments:
                    rels_tree = lxml_etree.fromstring(raw)
                    rels_ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
                    existing = [
                        r for r in rels_tree.findall(f'{{{rels_ns}}}Relationship')
                        if r.get('Type') == COMMENTS_REL_TYPE
                    ]
                    if not existing:
                        used_ids = {
                            r.get('Id')
                            for r in rels_tree.findall(f'{{{rels_ns}}}Relationship')
                        }
                        rid_num = 100
                        while f'rId{rid_num}' in used_ids:
                            rid_num += 1
                        lxml_etree.SubElement(rels_tree, f'{{{rels_ns}}}Relationship', {
                            'Id': f'rId{rid_num}',
                            'Type': COMMENTS_REL_TYPE,
                            'Target': 'comments.xml',
                        })
                    zout.writestr(item, lxml_etree.tostring(
                        rels_tree, xml_declaration=True, encoding="UTF-8", standalone=True
                    ))
                else:
                    zout.writestr(item, raw)

            if has_comments and not has_comments_xml:
                zout.writestr("word/comments.xml", comments_xml_bytes)

    os.replace(final_temp, output_path)

    # Clean up temp file
    if os.path.exists(temp_path):
        os.remove(temp_path)

    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="Generate an AIS-branded proposal Word document from JSON input."
    )
    parser.add_argument(
        "--input",
        help="Path to JSON input file conforming to proposal-schema.json.",
    )
    parser.add_argument(
        "--output", help="Path for the output .docx file."
    )
    parser.add_argument(
        "--template",
        default=str(DEFAULT_TEMPLATE),
        help="Path to the template .docx file. Defaults to assets/template.docx.",
    )

    parser.add_argument(
        "--validate",
        nargs="?",
        const="AUTO",
        metavar="FILE",
        help=(
            "Validate a .docx file without generating. "
            "If used with --input/--output, validates the output after generation."
        ),
    )
    parser.add_argument(
        "--strict",
        action="store_true",
        help="Treat warnings as errors during validation.",
    )
    parser.add_argument(
        "--skip-validation",
        action="store_true",
        help="Skip post-generation validation.",
    )

    args = parser.parse_args()

    # Standalone validation mode
    if args.validate and args.validate != "AUTO":
        issues = validate_docx(args.validate, strict=args.strict)
        ok = print_validation_results(issues, args.validate)
        sys.exit(0 if ok else 1)

    if not args.input:
        parser.error("--input is required for generation")
    if not args.output:
        parser.error("--output is required for generation")

    # Validate input
    print(f"Loading and validating input: {args.input}")
    try:
        data = load_and_validate(args.input)
    except jsonschema.ValidationError as e:
        print(f"Validation error: {e.message}", file=sys.stderr)
        print(f"  Path: {' -> '.join(str(p) for p in e.absolute_path)}", file=sys.stderr)
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"JSON parse error: {e}", file=sys.stderr)
        sys.exit(1)

    print(f"Using template: {args.template}")
    if not os.path.exists(args.template):
        print(f"Error: Template not found: {args.template}", file=sys.stderr)
        sys.exit(1)

    # Generate
    output = process_template(args.template, data, args.output)
    print(f"Generated proposal: {output}")

    # Summary
    sections = data.get("sections", [])
    comments = data.get("comments", [])
    print(f"  Sections: {len(sections)}")
    print(f"  Comments: {len(comments)}")
    print(f"  Cover page: {data['cover_page']['title']}")

    # Post-generation validation
    if not args.skip_validation:
        print()
        issues = validate_docx(output, strict=args.strict)
        ok = print_validation_results(issues, output)
        if not ok:
            print("\nGenerated file has errors. Use --skip-validation to bypass.", file=sys.stderr)
            sys.exit(1)


if __name__ == "__main__":
    main()
