# /// script
# dependencies = [
#   "python-docx>=1.1.0",
#   "jsonschema>=4.20.0",
# ]
# requires-python = ">=3.10"
# ///

"""
AIS Branded Document Generator

Reads a JSON input file, validates it against the document schema, opens the
AIS-branded template.docx, performs placeholder replacement while preserving
all existing styles, generates body content sections, and writes the output
.docx file.

Usage:
    uv run scripts/generate.py --input data.json --output document.docx
    uv run scripts/generate.py --validate document.docx
"""

import argparse
import json
import os
import re
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from xml.etree import ElementTree as ET

import jsonschema
from docx import Document
from docx.oxml.ns import qn
from lxml import etree as lxml_etree

SKILL_DIR = Path(__file__).resolve().parent.parent
DEFAULT_TEMPLATE = SKILL_DIR / "assets" / "template.docx"
SCHEMA_PATH = SKILL_DIR / "assets" / "document-schema.json"

NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Maps content block types from JSON schema to Word style names
CONTENT_STYLE_MAP = {
    "body": "Normal",
    "first_paragraph": "FirstParagraph",
    "heading2": "Heading2",
    "heading3": "Heading3",
    "heading4": "Heading4",
    "heading5": "Heading5",
    "compact": "Compact",
    "list_paragraph": "ListParagraph",
    "bullet": "ListParagraph",
    "numbered_list": "ListParagraph",
    "body_text": "BodyText",
    "source_code": "SourceCode",
    "bold_blue": "AISBoldBlue",
}

# Numbering IDs from the template's numbering.xml
BULLET_NUM_ID = "26"   # abstractNum 0 → bullet (•)
DECIMAL_NUM_ID = "29"  # abstractNum 2 → decimal (1. 2. 3.)

# Content types that get bullet or number formatting
BULLET_TYPES = {"list_paragraph", "bullet"}
NUMBERED_TYPES = {"numbered_list"}


def load_and_validate(json_path: str) -> dict:
    """Load JSON input and validate against schema."""
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    with open(SCHEMA_PATH, "r", encoding="utf-8") as f:
        schema = json.load(f)

    jsonschema.validate(instance=data, schema=schema)
    return data


def build_replacements(data: dict) -> dict:
    """Build a map of placeholder → replacement value from input data."""
    cover = data.get("cover_page", {})
    footer = data.get("footer", {})

    date_str = cover.get("date", datetime.now().strftime("%-m/%-d/%Y"))

    # Parse the display date into ISO format for the w:fullDate attribute
    full_date_iso = _parse_date_to_iso(date_str)

    return {
        "{{document_title}}": cover.get("title", ""),
        "{{subtitle}}": cover.get("subtitle", ""),
        "{{date}}": date_str,
        "{{full_date}}": full_date_iso,
        "{{prepared_for}}": cover.get("prepared_for", ""),
        "{{prepared_by}}": cover.get("prepared_by", "Applied Information Sciences, Inc."),
        "{{version}}": footer.get("version", cover.get("version", "v1.0")),
    }


def _parse_date_to_iso(date_str: str) -> str:
    """Convert a display date string to ISO 8601 for w:fullDate attribute."""
    for fmt in ("%m/%d/%Y", "%-m/%-d/%Y", "%m/%d/%y", "%Y-%m-%d", "%B %d, %Y"):
        try:
            dt = datetime.strptime(date_str, fmt)
            return dt.strftime("%Y-%m-%dT00:00:00Z")
        except ValueError:
            continue
    return datetime.now().strftime("%Y-%m-%dT00:00:00Z")


def _replace_in_xml_part(xml_bytes: bytes, replacements: dict) -> bytes:
    """Replace placeholder strings in an XML part's text content.

    Values are XML-escaped to prevent parse errors from characters like & < >.
    """
    from xml.sax.saxutils import escape

    content = xml_bytes.decode("utf-8")
    for placeholder, value in replacements.items():
        content = content.replace(placeholder, escape(str(value)))
    return content.encode("utf-8")


def _fix_mc_ignorable_namespaces(doc_xml_bytes: bytes, ns_map: dict) -> bytes:
    """Ensure every prefix listed in mc:Ignorable has a matching xmlns declaration."""
    doc_str = doc_xml_bytes.decode("utf-8")

    ignorable_match = re.search(r'mc:Ignorable="([^"]*)"', doc_str)
    if not ignorable_match:
        return doc_xml_bytes

    prefixes = ignorable_match.group(1).split()
    insertions = []
    for prefix in prefixes:
        if prefix in ns_map and f"xmlns:{prefix}=" not in doc_str:
            insertions.append(f'xmlns:{prefix}="{ns_map[prefix]}"')

    if not insertions:
        return doc_xml_bytes

    insert_str = " " + " ".join(insertions)
    doc_str = doc_str.replace(
        f'mc:Ignorable="{ignorable_match.group(1)}"',
        insert_str + f' mc:Ignorable="{ignorable_match.group(1)}"',
    )

    return doc_str.encode("utf-8")


def process_template(template_path: str, data: dict, output_path: str):
    """Main processing pipeline: template → filled document."""
    replacements = build_replacements(data)

    # Step 1: Raw XML replacement for headers/footers/cover page
    import tempfile

    temp_path = output_path + ".tmp"

    xml_parts_to_replace = [
        "word/document.xml",
        "word/header1.xml", "word/header2.xml", "word/header3.xml",
        "word/header4.xml",
        "word/footer1.xml", "word/footer2.xml", "word/footer3.xml",
        "word/footer4.xml", "word/footer5.xml",
    ]

    with zipfile.ZipFile(template_path, "r") as zin:
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                raw = zin.read(item.filename)
                if item.filename in xml_parts_to_replace:
                    raw = _replace_in_xml_part(raw, replacements)
                zout.writestr(item, raw)

    # Step 2: Open with python-docx for body content manipulation
    doc = Document(temp_path)

    # Step 3: Find and clear template body paragraphs (between cover and end)
    body_elem = doc.element.body
    paragraphs = body_elem.findall(qn("w:p"))

    # Find the TOC placeholder paragraph and the first Heading1 placeholder
    toc_placeholder_para = None
    remove_start = None
    for i, para in enumerate(paragraphs):
        text = "".join(t.text or "" for t in para.findall(f".//{qn('w:t')}"))
        if "{{toc_placeholder}}" in text:
            toc_placeholder_para = para
        ppr = para.find(qn("w:pPr"))
        if ppr is not None:
            style_el = ppr.find(qn("w:pStyle"))
            if style_el is not None:
                style_val = style_el.get(qn("w:val"), "")
                if style_val == "Heading1" and "{{section_heading}}" in text:
                    remove_start = i
                    break

    if remove_start is not None:
        for para in paragraphs[remove_start:]:
            body_elem.remove(para)

    # Step 3b: Replace TOC placeholder with a real TOC field + static entries
    sections = data.get("sections", [])
    if toc_placeholder_para is not None:
        _build_toc(body_elem, toc_placeholder_para, sections)

    # Step 4: Add sections from input data
    sect_pr = body_elem.find(qn("w:sectPr"))

    for section in sections:
        # Add Heading1
        heading_para = lxml_etree.SubElement(body_elem, qn("w:p"))
        h_ppr = lxml_etree.SubElement(heading_para, qn("w:pPr"))
        h_style = lxml_etree.SubElement(h_ppr, qn("w:pStyle"))
        h_style.set(qn("w:val"), "Heading1")
        h_run = lxml_etree.SubElement(heading_para, qn("w:r"))
        h_text = lxml_etree.SubElement(h_run, qn("w:t"))
        h_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        h_text.text = section["heading"]

        # Add content blocks
        for block in section.get("content", []):
            block_type = block.get("type", "body")

            # Handle table blocks separately
            if block_type == "table":
                _add_table(body_elem, block)
                continue

            style_name = CONTENT_STYLE_MAP.get(block_type, "Normal")
            text = block.get("text", "")
            bold_label = block.get("bold_label")

            para = lxml_etree.SubElement(body_elem, qn("w:p"))
            ppr = lxml_etree.SubElement(para, qn("w:pPr"))
            ps = lxml_etree.SubElement(ppr, qn("w:pStyle"))
            ps.set(qn("w:val"), style_name)

            # Add bullet or number formatting
            if block_type in BULLET_TYPES:
                num_pr = lxml_etree.SubElement(ppr, qn("w:numPr"))
                ilvl = lxml_etree.SubElement(num_pr, qn("w:ilvl"))
                ilvl.set(qn("w:val"), str(block.get("level", 0)))
                num_id = lxml_etree.SubElement(num_pr, qn("w:numId"))
                num_id.set(qn("w:val"), BULLET_NUM_ID)
            elif block_type in NUMBERED_TYPES:
                num_pr = lxml_etree.SubElement(ppr, qn("w:numPr"))
                ilvl = lxml_etree.SubElement(num_pr, qn("w:ilvl"))
                ilvl.set(qn("w:val"), str(block.get("level", 0)))
                num_id = lxml_etree.SubElement(num_pr, qn("w:numId"))
                num_id.set(qn("w:val"), DECIMAL_NUM_ID)

            if bold_label:
                label_run = lxml_etree.SubElement(para, qn("w:r"))
                label_rpr = lxml_etree.SubElement(label_run, qn("w:rPr"))
                lxml_etree.SubElement(label_rpr, qn("w:b"))
                label_t = lxml_etree.SubElement(label_run, qn("w:t"))
                label_t.set(
                    "{http://www.w3.org/XML/1998/namespace}space", "preserve"
                )
                label_t.text = bold_label + " "

            if block_type == "source_code":
                _add_source_code_runs(para, text)
            else:
                _add_formatted_runs(para, text)

        # Empty paragraph after section
        empty = lxml_etree.SubElement(body_elem, qn("w:p"))

    # Move sectPr to end
    if sect_pr is not None:
        body_elem.remove(sect_pr)
        body_elem.append(sect_pr)

    # Step 5: Save intermediate
    doc.save(output_path)

    # Step 6: Post-process - fix namespace declarations in all XML parts
    MC_IGNORABLE_NAMESPACES = {
        "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
        "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
        "w16": "http://schemas.microsoft.com/office/word/2018/wordml",
        "w16cex": "http://schemas.microsoft.com/office/word/2018/wordml/cex",
        "w16cid": "http://schemas.microsoft.com/office/word/2016/wordml/cid",
        "w16du": "http://schemas.microsoft.com/office/word/2023/wordml/word16du",
        "w16sdtdh": "http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash",
        "w16sdtfl": "http://schemas.microsoft.com/office/word/2021/wordml/sdtformatlock",
        "w16se": "http://schemas.microsoft.com/office/word/2015/wordml/symex",
        "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing14",
    }

    comments_data = data.get("comments", [])
    has_comments = len(comments_data) > 0
    comments_xml_bytes = None

    if has_comments:
        # Build comments XML
        comments_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        comments_root = lxml_etree.Element(
            qn("w:comments"),
            nsmap={
                "w": comments_ns,
                "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            },
        )
        for i, comment in enumerate(comments_data):
            cid = str(200 + i)
            author = comment.get("author", "AIS Template")
            date_str = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
            c_elem = lxml_etree.SubElement(comments_root, qn("w:comment"))
            c_elem.set(qn("w:id"), cid)
            c_elem.set(qn("w:author"), author)
            c_elem.set(qn("w:date"), date_str)
            c_p = lxml_etree.SubElement(c_elem, qn("w:p"))
            c_r = lxml_etree.SubElement(c_p, qn("w:r"))
            c_t = lxml_etree.SubElement(c_r, qn("w:t"))
            c_t.text = comment["text"]

        comments_xml_bytes = lxml_etree.tostring(
            comments_root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    COMMENTS_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
    COMMENTS_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"

    final_temp = output_path + ".final.tmp"
    with zipfile.ZipFile(output_path, "r") as zin:
        with zipfile.ZipFile(final_temp, "w", zipfile.ZIP_DEFLATED) as zout:
            has_comments_xml = False
            for item in zin.infolist():
                raw = zin.read(item.filename)

                if item.filename == "word/comments.xml" and has_comments:
                    zout.writestr(item, comments_xml_bytes)
                    has_comments_xml = True
                elif item.filename.endswith(".xml") and b"mc:Ignorable" in raw:
                    raw = _fix_mc_ignorable_namespaces(raw, MC_IGNORABLE_NAMESPACES)
                    zout.writestr(item, raw)
                elif item.filename == "[Content_Types].xml" and has_comments:
                    ct_tree = lxml_etree.fromstring(raw)
                    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
                    existing = ct_tree.findall(
                        f'{{{ct_ns}}}Override[@PartName="/word/comments.xml"]'
                    )
                    if not existing:
                        lxml_etree.SubElement(
                            ct_tree,
                            f"{{{ct_ns}}}Override",
                            {
                                "PartName": "/word/comments.xml",
                                "ContentType": COMMENTS_CT,
                            },
                        )
                    zout.writestr(
                        item,
                        lxml_etree.tostring(
                            ct_tree,
                            xml_declaration=True,
                            encoding="UTF-8",
                            standalone=True,
                        ),
                    )
                elif item.filename == "word/_rels/document.xml.rels" and has_comments:
                    rels_tree = lxml_etree.fromstring(raw)
                    rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
                    existing = [
                        r
                        for r in rels_tree.findall(f"{{{rels_ns}}}Relationship")
                        if r.get("Type") == COMMENTS_REL_TYPE
                    ]
                    if not existing:
                        used_ids = {
                            r.get("Id")
                            for r in rels_tree.findall(f"{{{rels_ns}}}Relationship")
                        }
                        rid_num = 100
                        while f"rId{rid_num}" in used_ids:
                            rid_num += 1
                        lxml_etree.SubElement(
                            rels_tree,
                            f"{{{rels_ns}}}Relationship",
                            {
                                "Id": f"rId{rid_num}",
                                "Type": COMMENTS_REL_TYPE,
                                "Target": "comments.xml",
                            },
                        )
                    zout.writestr(
                        item,
                        lxml_etree.tostring(
                            rels_tree,
                            xml_declaration=True,
                            encoding="UTF-8",
                            standalone=True,
                        ),
                    )
                else:
                    zout.writestr(item, raw)

            if has_comments and not has_comments_xml:
                zout.writestr("word/comments.xml", comments_xml_bytes)

    os.replace(final_temp, output_path)

    # Clean up temp file
    if os.path.exists(temp_path):
        os.remove(temp_path)

    return output_path


def _build_toc(body_elem, toc_placeholder_para, sections: list):
    """Replace the TOC placeholder paragraph with a TOC field and static entries."""
    parent = toc_placeholder_para.getparent()
    insert_idx = list(parent).index(toc_placeholder_para)
    parent.remove(toc_placeholder_para)

    # Build TOC field: begin → instrText → separate → [static entries] → end
    # Field begin paragraph
    p_begin = lxml_etree.Element(qn("w:p"))
    ppr = lxml_etree.SubElement(p_begin, qn("w:pPr"))
    ps = lxml_etree.SubElement(ppr, qn("w:pStyle"))
    ps.set(qn("w:val"), "TOC1")

    r_begin = lxml_etree.SubElement(p_begin, qn("w:r"))
    fc_begin = lxml_etree.SubElement(r_begin, qn("w:fldChar"))
    fc_begin.set(qn("w:fldCharType"), "begin")

    r_instr = lxml_etree.SubElement(p_begin, qn("w:r"))
    instr_text = lxml_etree.SubElement(r_instr, qn("w:instrText"))
    instr_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '

    r_sep = lxml_etree.SubElement(p_begin, qn("w:r"))
    fc_sep = lxml_etree.SubElement(r_sep, qn("w:fldChar"))
    fc_sep.set(qn("w:fldCharType"), "separate")

    parent.insert(insert_idx, p_begin)
    insert_idx += 1

    # Static TOC entries for each section (so the TOC isn't empty before update)
    for section in sections:
        heading = section.get("heading", "")
        toc_entry = lxml_etree.Element(qn("w:p"))
        t_ppr = lxml_etree.SubElement(toc_entry, qn("w:pPr"))
        t_ps = lxml_etree.SubElement(t_ppr, qn("w:pStyle"))
        t_ps.set(qn("w:val"), "TOC1")
        t_run = lxml_etree.SubElement(toc_entry, qn("w:r"))
        t_text = lxml_etree.SubElement(t_run, qn("w:t"))
        t_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t_text.text = heading

        # Also add TOC2 entries for heading2 blocks
        parent.insert(insert_idx, toc_entry)
        insert_idx += 1

        for block in section.get("content", []):
            if block.get("type") in ("heading2", "heading3"):
                toc_level = "TOC2" if block["type"] == "heading2" else "TOC3"
                sub_entry = lxml_etree.Element(qn("w:p"))
                s_ppr = lxml_etree.SubElement(sub_entry, qn("w:pPr"))
                s_ps = lxml_etree.SubElement(s_ppr, qn("w:pStyle"))
                s_ps.set(qn("w:val"), toc_level)
                s_run = lxml_etree.SubElement(sub_entry, qn("w:r"))
                s_text = lxml_etree.SubElement(s_run, qn("w:t"))
                s_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                s_text.text = block.get("text", "")
                parent.insert(insert_idx, sub_entry)
                insert_idx += 1

    # Field end paragraph
    p_end = lxml_etree.Element(qn("w:p"))
    r_end = lxml_etree.SubElement(p_end, qn("w:r"))
    fc_end = lxml_etree.SubElement(r_end, qn("w:fldChar"))
    fc_end.set(qn("w:fldCharType"), "end")
    parent.insert(insert_idx, p_end)


def _add_table(body_elem, block: dict):
    """Add a Word table from a table content block."""
    rows_data = block.get("rows", [])
    if not rows_data:
        return

    has_header = block.get("header", True)

    tbl = lxml_etree.SubElement(body_elem, qn("w:tbl"))

    # Table properties
    tbl_pr = lxml_etree.SubElement(tbl, qn("w:tblPr"))
    tbl_style = lxml_etree.SubElement(tbl_pr, qn("w:tblStyle"))
    tbl_style.set(qn("w:val"), "GridTable4-Accent1")
    tbl_w = lxml_etree.SubElement(tbl_pr, qn("w:tblW"))
    tbl_w.set(qn("w:w"), "0")
    tbl_w.set(qn("w:type"), "auto")
    look = lxml_etree.SubElement(tbl_pr, qn("w:tblLook"))
    look.set(qn("w:val"), "04A0")
    look.set(qn("w:firstRow"), "1" if has_header else "0")
    look.set(qn("w:lastRow"), "0")
    look.set(qn("w:firstColumn"), "0")
    look.set(qn("w:lastColumn"), "0")
    look.set(qn("w:noHBand"), "0")
    look.set(qn("w:noVBand"), "1")

    # Table grid (auto-width columns)
    tbl_grid = lxml_etree.SubElement(tbl, qn("w:tblGrid"))
    if rows_data:
        num_cols = len(rows_data[0])
        for _ in range(num_cols):
            lxml_etree.SubElement(tbl_grid, qn("w:gridCol"))

    for row_data in rows_data:
        tr = lxml_etree.SubElement(tbl, qn("w:tr"))
        for cell_text in row_data:
            tc = lxml_etree.SubElement(tr, qn("w:tc"))
            p = lxml_etree.SubElement(tc, qn("w:p"))
            _add_formatted_runs(p, str(cell_text))


def _add_source_code_runs(para, text: str):
    """Add code text to a paragraph with proper line breaks via <w:br/> elements."""
    lines = text.split("\n")
    for idx, line in enumerate(lines):
        run = lxml_etree.SubElement(para, qn("w:r"))
        t = lxml_etree.SubElement(run, qn("w:t"))
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = line
        if idx < len(lines) - 1:
            br_run = lxml_etree.SubElement(para, qn("w:r"))
            lxml_etree.SubElement(br_run, qn("w:br"))


def _add_formatted_runs(para, text: str):
    """Add text runs to a paragraph, handling markdown-style bold/italic/code."""
    # Pattern: **bold**, *italic*, `code`
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)"
    parts = re.split(pattern, text)

    i = 0
    while i < len(parts):
        part = parts[i]
        if part is None:
            i += 1
            continue

        if i + 1 < len(parts) and parts[i] and parts[i].startswith("**"):
            # Bold
            bold_text = parts[i + 1]
            run = lxml_etree.SubElement(para, qn("w:r"))
            rpr = lxml_etree.SubElement(run, qn("w:rPr"))
            lxml_etree.SubElement(rpr, qn("w:b"))
            t = lxml_etree.SubElement(run, qn("w:t"))
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = bold_text
            i += 3
        elif i + 1 < len(parts) and parts[i] and parts[i].startswith("*") and not parts[i].startswith("**"):
            # Italic
            italic_text = parts[i + 1] if i + 1 < len(parts) else ""
            run = lxml_etree.SubElement(para, qn("w:r"))
            rpr = lxml_etree.SubElement(run, qn("w:rPr"))
            lxml_etree.SubElement(rpr, qn("w:i"))
            t = lxml_etree.SubElement(run, qn("w:t"))
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = italic_text
            i += 3
        elif i + 1 < len(parts) and parts[i] and parts[i].startswith("`"):
            # Code
            code_text = parts[i + 1] if i + 1 < len(parts) else ""
            run = lxml_etree.SubElement(para, qn("w:r"))
            rpr = lxml_etree.SubElement(run, qn("w:rPr"))
            rs = lxml_etree.SubElement(rpr, qn("w:rStyle"))
            rs.set(qn("w:val"), "VerbatimChar")
            t = lxml_etree.SubElement(run, qn("w:t"))
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = code_text
            i += 3
        elif part:
            # Plain text
            run = lxml_etree.SubElement(para, qn("w:r"))
            t = lxml_etree.SubElement(run, qn("w:t"))
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = part
            i += 1
        else:
            i += 1


def validate_docx(path: str, strict: bool = False) -> list[dict]:
    """Validate a .docx file for common OOXML issues that cause Word errors."""
    issues = []

    def _add(severity: str, part: str, msg: str):
        issues.append({"severity": severity, "part": part, "message": msg})

    if not os.path.exists(path):
        _add("error", path, "File does not exist")
        return issues

    try:
        zf = zipfile.ZipFile(path, "r")
    except zipfile.BadZipFile:
        _add("error", path, "Not a valid ZIP/DOCX file")
        return issues

    with zf:
        names = set(zf.namelist())

        for required in ["[Content_Types].xml", "word/document.xml"]:
            if required not in names:
                _add("error", required, "Required part is missing")

        for name in sorted(names):
            if not name.endswith(".xml") and not name.endswith(".rels"):
                continue
            try:
                raw = zf.read(name)
                ET.fromstring(raw)
            except ET.ParseError as e:
                _add("error", name, f"Malformed XML: {e}")

        mc_ns_map = {
            "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
            "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
            "w16": "http://schemas.microsoft.com/office/word/2018/wordml",
            "w16cex": "http://schemas.microsoft.com/office/word/2018/wordml/cex",
            "w16cid": "http://schemas.microsoft.com/office/word/2016/wordml/cid",
            "w16du": "http://schemas.microsoft.com/office/word/2023/wordml/word16du",
            "w16sdtdh": "http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash",
            "w16sdtfl": "http://schemas.microsoft.com/office/word/2021/wordml/sdtformatlock",
            "w16se": "http://schemas.microsoft.com/office/word/2015/wordml/symex",
            "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing14",
        }
        for name in sorted(names):
            if not name.endswith(".xml"):
                continue
            content = zf.read(name).decode("utf-8", errors="replace")
            ig_match = re.search(r'mc:Ignorable="([^"]*)"', content)
            if not ig_match:
                continue
            for prefix in ig_match.group(1).split():
                if f"xmlns:{prefix}=" not in content:
                    _add(
                        "error",
                        name,
                        f'mc:Ignorable lists "{prefix}" but xmlns:{prefix} is not declared',
                    )

        if "word/document.xml" in names:
            doc_content = zf.read("word/document.xml").decode("utf-8", errors="replace")
            comment_refs = set(
                re.findall(r'w:commentRangeStart\s[^>]*w:id="(\d+)"', doc_content)
            )
            if comment_refs and "word/comments.xml" not in names:
                _add(
                    "error",
                    "word/comments.xml",
                    f"Document has {len(comment_refs)} comment anchors but no comments.xml",
                )

        sensitive_patterns = [
            (r"[A-Z]:\\Users\\[^\"<]+", "Local file path detected"),
            (r"file:///[A-Z]:", "file:// URI with drive letter detected"),
        ]
        skip_parts = {"word/fontTable.xml"}
        for name in sorted(names):
            if not name.endswith(".xml") or name in skip_parts:
                continue
            content = zf.read(name).decode("utf-8", errors="replace")
            for pattern, msg in sensitive_patterns:
                matches = re.findall(pattern, content)
                if matches:
                    _add("warning", name, f"{msg}: {matches[0]}")

    if strict:
        for issue in issues:
            if issue["severity"] == "warning":
                issue["severity"] = "error"

    return issues


def print_validation_results(issues: list[dict], path: str) -> bool:
    """Print validation results. Returns True if no errors."""
    errors = [i for i in issues if i["severity"] == "error"]
    warnings = [i for i in issues if i["severity"] == "warning"]

    if not issues:
        print(f"✅ {path}: All checks passed")
        return True

    if errors:
        print(f"❌ {path}: {len(errors)} error(s), {len(warnings)} warning(s)")
    else:
        print(f"⚠️  {path}: {len(warnings)} warning(s)")

    for issue in issues:
        icon = "❌" if issue["severity"] == "error" else "⚠️ "
        print(f"  {icon} [{issue['part']}] {issue['message']}")

    return len(errors) == 0


def main():
    parser = argparse.ArgumentParser(
        description="Generate an AIS-branded Word document from JSON input."
    )
    parser.add_argument(
        "--input",
        help="Path to JSON input file conforming to document-schema.json.",
    )
    parser.add_argument("--output", help="Path for the output .docx file.")
    parser.add_argument(
        "--template",
        default=str(DEFAULT_TEMPLATE),
        help="Path to the template .docx file. Defaults to assets/template.docx.",
    )
    parser.add_argument(
        "--validate",
        nargs="?",
        const="AUTO",
        metavar="FILE",
        help="Validate a .docx file without generating.",
    )
    parser.add_argument(
        "--strict",
        action="store_true",
        help="Treat warnings as errors during validation.",
    )
    parser.add_argument(
        "--skip-validation",
        action="store_true",
        help="Skip post-generation validation.",
    )

    args = parser.parse_args()

    # Standalone validation mode
    if args.validate and args.validate != "AUTO":
        issues = validate_docx(args.validate, strict=args.strict)
        ok = print_validation_results(issues, args.validate)
        sys.exit(0 if ok else 1)

    if not args.input:
        parser.error("--input is required for generation")
    if not args.output:
        parser.error("--output is required for generation")

    # Validate input
    print(f"Loading and validating input: {args.input}")
    try:
        data = load_and_validate(args.input)
    except jsonschema.ValidationError as e:
        print(f"Validation error: {e.message}", file=sys.stderr)
        print(
            f"  Path: {' -> '.join(str(p) for p in e.absolute_path)}", file=sys.stderr
        )
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"JSON parse error: {e}", file=sys.stderr)
        sys.exit(1)

    print(f"Using template: {args.template}")
    if not os.path.exists(args.template):
        print(f"Error: Template not found: {args.template}", file=sys.stderr)
        sys.exit(1)

    # Generate
    output = process_template(args.template, data, args.output)
    print(f"Generated document: {output}")

    # Summary
    sections = data.get("sections", [])
    comments = data.get("comments", [])
    print(f"  Sections: {len(sections)}")
    print(f"  Comments: {len(comments)}")
    print(f"  Title: {data['cover_page']['title']}")

    # Post-generation validation
    if not args.skip_validation:
        print()
        issues = validate_docx(output, strict=args.strict)
        ok = print_validation_results(issues, output)
        if not ok:
            print(
                "\nGenerated file has errors. Use --skip-validation to bypass.",
                file=sys.stderr,
            )
            sys.exit(1)


if __name__ == "__main__":
    main()
